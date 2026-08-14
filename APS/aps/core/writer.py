#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""文字引擎：docx / txt 读写。

模型：段落列表 [{text, style, bold, size, color, align, list}]
与 python-docx 互转，供 GTK4 富文本视图使用。
"""
import os

from docx import Document as DocxDoc
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _rgb(hex_str):
    hex_str = hex_str.lstrip("#")
    try:
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))
    except Exception:
        return None


class WriterEngine:
    def __init__(self, path: str = None, text: str = None):
        self.path = path
        self.paragraphs = []      # [{text, style, bold, italic, size, color, align, list}]
        self.dirty = False
        if text is not None:
            for line in text.splitlines():
                self.paragraphs.append(_para(line))
        elif path:
            self._load(path)

    # ---------------- 模型 ----------------
    def new_para(self, text="", **kw):
        p = _para(text, **kw)
        self.paragraphs.append(p)
        self.dirty = True
        return p

    def clear(self):
        self.paragraphs = []
        self.dirty = True

    # ---------------- 读写 ----------------
    def _load(self, path: str):
        doc = DocxDoc(path)
        self.paragraphs = []
        for p in doc.paragraphs:
            runs = p.runs
            bold = any(r.bold for r in runs) if runs else None
            italic = any(r.italic for r in runs) if runs else None
            size = None
            color = None
            for r in runs:
                if r.font.size:
                    size = r.font.size.pt
                if r.font.color and r.font.color.rgb:
                    color = str(r.font.color.rgb)
                if size and color:
                    break
            self.paragraphs.append(_para(
                p.text, bold=bold, italic=italic, size=size, color=color,
                align=_align_str(p.alignment),
            ))
        self.dirty = False

    def save(self, path: str):
        doc = DocxDoc()
        for p in self.paragraphs:
            para = doc.add_paragraph()
            if p.get("style") == "title":
                para.style = doc.styles["Title"]
            elif p.get("style") == "heading":
                para.style = doc.styles["Heading 1"]
            run = para.add_run(p.get("text", ""))
            if p.get("bold"):
                run.bold = True
            if p.get("italic"):
                run.italic = True
            if p.get("size"):
                run.font.size = Pt(float(p["size"]))
            if p.get("color"):
                rgb = _rgb(p["color"])
                if rgb:
                    run.font.color.rgb = rgb
            align = p.get("align")
            if align:
                para.alignment = _align_enum(align)
        doc.save(path)
        self.path = path
        self.dirty = False

    def to_text(self) -> str:
        return "\n".join(p.get("text", "") for p in self.paragraphs)


# ---------------- 辅助 ----------------
def _para(text="", style=None, bold=None, italic=None, size=None, color=None, align=None):
    return {"text": text, "style": style, "bold": bold, "italic": italic,
            "size": size, "color": color, "align": align}


def _align_str(a):
    if a is None:
        return None
    m = {0: "left", 1: "center", 2: "right", 3: "justify"}
    return m.get(int(a))


def _align_enum(s):
    m = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
         "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    return m.get(s)
