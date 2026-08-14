import gi
gi.require_version("Gtk", "4.0")
from gi.repository import Gtk
from pathlib import Path


class Document:
    FORMATS = {}
    extensions = ()

    def read(self):
        raise NotImplementedError

    def write(self, text):
        raise NotImplementedError

    def get_display_name(self):
        raise NotImplementedError

    def __init_subclass__(cls, **kwargs):
        super().__init_subclass__(**kwargs)
        for ext in cls.extensions:
            Document.FORMATS[ext] = cls


class MdDocument(Document):
    extensions = (".md", ".txt", ".markdown")

    def __init__(self, path):
        self.path = Path(path)

    def read(self):
        return self.path.read_text(encoding="utf-8")

    def write(self, text):
        self.path.write_text(text, encoding="utf-8")

    def get_display_name(self):
        return "Markdown"


class TextDocument(Document):
    extensions = ()

    def __init__(self, path):
        self.path = Path(path)

    def read(self):
        return self.path.read_text(encoding="utf-8", errors="replace")

    def write(self, text):
        self.path.write_text(text, encoding="utf-8")

    def get_display_name(self):
        return "Text"


class DocxDocument(Document):
    extensions = (".docx",)

    def __init__(self, path):
        self.path = Path(path)

    def read(self):
        from docx import Document as DocxDoc
        doc = DocxDoc(str(self.path))
        paragraphs = []
        for p in doc.paragraphs:
            paragraphs.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    def write(self, text):
        from docx import Document as DocxDoc
        doc = DocxDoc()
        for line in text.split("\n"):
            if line.strip().startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                heading = line.lstrip("#").strip()
                doc.add_heading(heading, level=min(level, 4))
            else:
                doc.add_paragraph(line)
        doc.save(str(self.path))

    def get_display_name(self):
        return "Word"


class XlsxDocument(Document):
    extensions = (".xlsx", ".xls")

    def __init__(self, path):
        self.path = Path(path)

    def read(self):
        import openpyxl
        wb = openpyxl.load_workbook(str(self.path), data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"## Sheet: {sheet_name}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                lines.append("\t".join(cells))
            lines.append("")
        wb.close()
        return "\n".join(lines)

    def write(self, text):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        current_sheet = ws
        row_idx = 1
        for line in text.split("\n"):
            if line.startswith("## Sheet: "):
                sheet_name = line[10:].strip()
                if sheet_name != "Sheet1" and sheet_name:
                    if row_idx > 1:
                        current_sheet = wb.create_sheet(title=sheet_name)
                        row_idx = 1
                continue
            if "\t" in line:
                cells = line.split("\t")
                for col_idx, val in enumerate(cells, 1):
                    current_sheet.cell(row=row_idx, column=col_idx, value=val)
                row_idx += 1
            elif line.strip() == "":
                row_idx += 1
        wb.save(str(self.path))

    def get_display_name(self):
        return "Excel"


class EditorPane(Gtk.Box):
    def __init__(self):
        super().__init__(orientation=Gtk.Orientation.VERTICAL, spacing=0)
        self.document = None
        self.filepath = None
        self._modified = False
        self._build_ui()

    def _build_ui(self):
        self._header = Gtk.Box(orientation=Gtk.Orientation.HORIZONTAL, spacing=6)
        self._header.set_margin_start(8)
        self._header.set_margin_end(8)
        self._header.set_margin_top(4)
        self._header.set_margin_bottom(4)

        self._title_label = Gtk.Label(label="未命名")
        self._title_label.add_css_class("heading")
        self._title_label.set_hexpand(True)
        self._title_label.set_xalign(0)

        self._format_label = Gtk.Label(label="")
        self._format_label.add_css_class("dim-label")

        self._header.append(self._title_label)
        self._header.append(self._format_label)

        self._scrolled = Gtk.ScrolledWindow()
        self._scrolled.set_vexpand(True)

        self._textview = Gtk.TextView()
        self._textview.set_wrap_mode(Gtk.WrapMode.WORD_CHAR)
        self._textview.set_monospace(False)
        self._textview.set_hexpand(True)
        self._textview.set_vexpand(True)
        self._textview.set_left_margin(12)
        self._textview.set_right_margin(12)
        self._textview.set_top_margin(8)
        self._textview.set_bottom_margin(8)



        self._buffer = self._textview.get_buffer()
        self._buffer.connect("changed", self._on_buffer_changed)

        self._scrolled.set_child(self._textview)

        self.append(self._header)
        self.append(self._scrolled)

    def _on_buffer_changed(self, *args):
        self._modified = True

    def open_file(self, path):
        path = Path(path)
        self.filepath = path
        ext = path.suffix.lower()
        doc_cls = Document.FORMATS.get(ext, TextDocument)
        self.document = doc_cls(path)
        text = self.document.read()
        self._buffer.set_text(text)
        self._modified = False
        self._title_label.set_label(path.name)
        self._format_label.set_label(self.document.get_display_name())
        return True

    def save(self):
        if not self.filepath or not self.document:
            return False
        bounds = self._buffer.get_bounds()
        text = self._buffer.get_text(bounds[0], bounds[1], False)
        self.document.write(text)
        self._modified = False
        return True

    def save_as(self, path):
        path = Path(path)
        ext = path.suffix.lower()
        doc_cls = Document.FORMATS.get(ext, TextDocument)
        bounds = self._buffer.get_bounds()
        text = self._buffer.get_text(bounds[0], bounds[1], False)
        doc = doc_cls(path)
        doc.write(text)
        self.filepath = path
        self.document = doc
        self._modified = False
        self._title_label.set_label(path.name)
        self._format_label.set_label(doc.get_display_name())
        return True

    def get_selected_text(self):
        bounds = self._buffer.get_selection_bounds()
        if bounds:
            return self._buffer.get_text(bounds[0], bounds[1], False)
        return ""

    def replace_selection(self, text):
        bounds = self._buffer.get_selection_bounds()
        if bounds:
            self._buffer.delete(bounds[0], bounds[1])
        self._buffer.insert_at_cursor(text)

    def insert_at_cursor(self, text):
        self._buffer.insert_at_cursor(text)

    def get_all_text(self):
        bounds = self._buffer.get_bounds()
        return self._buffer.get_text(bounds[0], bounds[1], False)

    @property
    def is_modified(self):
        return self._modified

    @property
    def is_empty(self):
        bounds = self._buffer.get_bounds()
        return self._buffer.get_text(bounds[0], bounds[1], False).strip() == ""
