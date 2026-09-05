import json
import os
import re
import sys

sys.path.insert(0, "/usr/chindows")

from flask import Flask, jsonify, render_template, request

from chindshell import chat as chatmod
from chindshell import flask as csf

app = Flask(__name__)
csf.register(app)

HOME = os.path.expanduser("~")
DEFAULT_DIR = os.path.join(HOME, "文档")
HISTORY_PATH = os.path.join(HOME, ".ainote2", "history.json")

state = {
    "dir": DEFAULT_DIR if os.path.isdir(DEFAULT_DIR) else HOME,
    "sid": None,
}

session = chatmod.StreamSession()
tracker = chatmod.ConversationTracker(HISTORY_PATH, session_provider=lambda: state.get("sid"))

EXT_ICON = {
    ".md": "📝", ".docx": "📄", ".xlsx": "📊", ".txt": "📃", ".json": "📋",
    ".yaml": "📋", ".yml": "📋", ".xml": "📋", ".html": "🌐", ".css": "🎨",
    ".js": "📜", ".py": "🐍", ".c": "⚙️", ".cpp": "⚙️", ".h": "⚙️",
    ".java": "☕", ".sh": "💻", ".csv": "📊",
}


# ---------------------------------------------------------------- 文件读写

def read_document(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        from docx import Document as DocxDoc
        doc = DocxDoc(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext in (".xlsx", ".xls"):
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        lines = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f"## Sheet: {sheet}")
            for row in ws.iter_rows(values_only=True):
                lines.append("\t".join("" if c is None else str(c) for c in row))
            lines.append("")
        wb.close()
        return "\n".join(lines)
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def write_document(path, text):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        from docx import Document as DocxDoc
        from docx.shared import Pt
        doc = DocxDoc()
        for line in text.split("\n"):
            if line.strip().startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                doc.add_heading(line.lstrip("#").strip(), level=min(level, 4))
            else:
                doc.add_paragraph(line)
        doc.save(path)
    elif ext in (".xlsx", ".xls"):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        row_idx = 1
        for line in text.split("\n"):
            if line.startswith("## Sheet: "):
                name = line[10:].strip()
                if name and name != "Sheet1" and row_idx > 1:
                    ws = wb.create_sheet(title=name)
                    row_idx = 1
                continue
            if "\t" in line:
                for col, val in enumerate(line.split("\t"), 1):
                    ws.cell(row=row_idx, column=col, value=val)
                row_idx += 1
            elif line.strip() == "":
                row_idx += 1
        wb.save(path)
    else:
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)


# ---------------------------------------------------------------- 页面 / 文件

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/dir")
def api_dir():
    names = []
    try:
        for f in sorted(os.scandir(state["dir"]), key=lambda e: (not e.is_file(), e.name.lower())):
            if f.is_file():
                names.append({"name": f.name, "ext": os.path.splitext(f.name)[1].lower()})
    except OSError:
        pass
    return jsonify({"dir": state["dir"], "files": names})


@app.route("/api/dir", methods=["POST"])
def api_set_dir():
    d = request.get_json(force=True) or {}
    p = os.path.expanduser((d.get("dir") or "").strip())
    if os.path.isdir(p):
        state["dir"] = p
        return jsonify({"ok": True, "dir": p})
    return jsonify({"ok": False, "error": "目录不存在"})


@app.route("/api/file")
def api_read_file():
    p = request.args.get("path", "")
    if not p or not os.path.isfile(p):
        return jsonify({"ok": False, "error": "文件不存在"})
    try:
        content = read_document(p)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})
    return jsonify({"ok": True, "path": p, "name": os.path.basename(p),
                    "ext": os.path.splitext(p)[1].lower(), "content": content})


@app.route("/api/file", methods=["POST"])
def api_write_file():
    d = request.get_json(force=True) or {}
    p = d.get("path", "")
    if not p:
        return jsonify({"ok": False, "error": "无文件路径"})
    try:
        write_document(p, d.get("content") or "")
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})
    return jsonify({"ok": True})


# ---------------------------------------------------------------- AI 会话

def _opencode_filter(line):
    try:
        ev = json.loads(line)
        t = ev.get("type")
        if t == "text":
            txt = ev.get("part", {}).get("text", "")
            if txt:
                return txt
        elif t in ("step_start", "step_finish"):
            sid = ev.get("sessionID")
            if sid:
                state["sid"] = sid
    except Exception:
        pass
    return None


@app.route("/api/ai/send", methods=["POST"])
def ai_send():
    d = request.get_json(force=True) or {}
    msg = (d.get("message") or "").strip()
    if not msg:
        return jsonify({"ok": False, "error": "空消息"})
    if session.running:
        return jsonify({"ok": False, "error": "AI 正在生成中"})
    mode = d.get("mode", "chat")
    full_doc = (d.get("full_doc") or "").strip()
    if full_doc:
        if mode == "selection":
            prompt = (f"当前文档内容：\n---\n{full_doc}\n---\n\n"
                      f"用户选中了以下文本，请按要求处理并只输出处理后的文本（不要包含文档其他部分）：\n{msg}")
        else:
            prompt = (f"当前文档内容：\n---\n{full_doc}\n---\n\n用户指令：{msg}\n\n"
                      "请根据指令输出完整的文档内容（包含所有修改后的完整文档，不要只修改输出部分）。")
    else:
        prompt = msg
    cmd = ["opencode", "run", "--format", "json"]
    if state.get("sid"):
        cmd += ["--session", state["sid"]]
    cmd.append(prompt)
    tracker.add_user(msg)
    session.start(cmd, line_filter=_opencode_filter)
    return jsonify({"ok": True})


@app.route("/api/ai/stream")
def ai_stream():
    since = int(request.args.get("since", 0) or 0)
    new, done, error, total = session.poll(since)
    if done:
        tracker.record_reply(session.full_text())
    return jsonify({"chunks": new, "done": done, "error": error, "total": total})


@app.route("/api/ai/stop")
def ai_stop():
    session.stop()
    return jsonify({"ok": True})


@app.route("/api/ai/new", methods=["POST"])
def ai_new():
    tracker.save()
    state["sid"] = None
    return jsonify({"ok": True})


# ---------------------------------------------------------------- 历史

@app.route("/api/history")
def api_history():
    return jsonify(tracker.load())


@app.route("/api/history/switch", methods=["POST"])
def api_history_switch():
    d = request.get_json(force=True) or {}
    session = d.get("session")
    msgs = d.get("messages") or []
    if session:
        state["sid"] = session  # 后续 opencode 用 --session 继续该会话
    tracker.messages = [{"role": m.get("role", "user"), "content": m.get("content", "")}
                        for m in msgs if (m.get("content") or "").strip()]
    tracker._assistant_pending = False
    tracker._live_index = -1
    return jsonify({"ok": True})


@app.route("/api/history/clear", methods=["POST"])
def api_history_clear():
    tracker.clear()
    return jsonify({"ok": True})
